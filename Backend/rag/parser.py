import logging
from pathlib import Path
from typing import Dict, Any

from rag.ocr import extract_ocr_from_pdf, extract_features_from_image

logger = logging.getLogger(__name__)


class DocumentParser:
    """
    Parses PDF, DOCX, TXT, and Image files (PNG, JPG, WEBP, etc.) to extract text, OCR, and visual features.
    """

    @staticmethod
    def parse(file_path: str) -> Dict[str, Any]:
        path = Path(file_path)
        ext = path.suffix.lower()

        if ext == ".pdf":
            return DocumentParser._parse_pdf(file_path)
        elif ext == ".docx":
            return DocumentParser._parse_docx(file_path)
        elif ext in [".txt", ".md", ".csv"]:
            return DocumentParser._parse_txt(file_path)
        elif ext in [".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff"]:
            return extract_features_from_image(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

    @staticmethod
    def _parse_pdf(file_path: str) -> Dict[str, Any]:
        extracted_pages = []
        full_text = ""

        # Step 1: Try PyMuPDF (fitz) for embedded text and page images
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(file_path)
            for page_num, page in enumerate(doc):
                text = page.get_text("text")
                img_count = len(page.get_images())
                
                if text.strip():
                    p_text = text
                    if img_count > 0:
                        p_text += f"\n[Page Features: {img_count} embedded figures/diagrams extracted]"
                    extracted_pages.append({
                        "page": page_num + 1,
                        "text": p_text
                    })
                    full_text += f"\n--- Page {page_num + 1} ---\n" + p_text
        except Exception as e:
            logger.warning(f"PyMuPDF text extraction failed on {file_path}: {e}")

        # Step 2: Fallback to pypdf if empty
        if not full_text.strip():
            try:
                import pypdf
                reader = pypdf.PdfReader(file_path)
                for page_num, page in enumerate(reader.pages):
                    text = page.extract_text() or ""
                    if text.strip():
                        extracted_pages.append({
                            "page": page_num + 1,
                            "text": text
                        })
                        full_text += f"\n--- Page {page_num + 1} ---\n" + text
            except Exception as e:
                logger.warning(f"pypdf failed on {file_path}: {e}")

        # Step 3: OCR and visual pixmap feature extraction for scanned / image PDFs
        if not full_text.strip() or len(full_text.strip()) < 50:
            logger.info(f"PDF {file_path} has missing or low text density. Running high-resolution pixmap OCR feature extraction...")
            ocr_result = extract_ocr_from_pdf(file_path)
            if ocr_result and ocr_result.get("text"):
                full_text = ocr_result["text"]
                extracted_pages = ocr_result.get("pages", [])

        return {
            "text": full_text.strip(),
            "pages": extracted_pages,
            "metadata": {
                "total_pages": len(extracted_pages),
                "is_ocr": not extracted_pages or "OCR" in full_text or "Scanned" in full_text
            }
        }

    @staticmethod
    def _parse_docx(file_path: str) -> Dict[str, Any]:
        try:
            import docx
            doc = docx.Document(file_path)
            full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            return {
                "text": full_text,
                "pages": [{"page": 1, "text": full_text}],
                "metadata": {"paragraphs": len(doc.paragraphs)}
            }
        except Exception as e:
            logger.error(f"Failed to parse DOCX {file_path}: {e}")
            raise

    @staticmethod
    def _parse_txt(file_path: str) -> Dict[str, Any]:
        text = Path(file_path).read_text(encoding="utf-8", errors="ignore")
        return {
            "text": text,
            "pages": [{"page": 1, "text": text}],
            "metadata": {"char_count": len(text)}
        }
